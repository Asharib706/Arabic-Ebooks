import streamlit as st
from pymongo import MongoClient
from bs4 import BeautifulSoup
from backend import TextToSpeech, html_to_docx
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Set page config
st.set_page_config(layout="wide", page_title="Arabic Ebook Viewer")
tts = TextToSpeech()

# Custom CSS with theme awareness
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Amiri:wght@400;700&display=swap');
            
:root {
    --text-color: #000000;
    --background-color: #ffffff;
    --secondary-background: #fafafa;
    --border-color: #e0e0e0;
    --shadow-color: rgba(0,0,0,0.05);
    --title-color: #2c3e50;
    --chapter-border: #e0e0e0;
}

[data-theme="dark"] {
    --text-color: #ffffff;
    --background-color: #0e1117;
    --secondary-background: #1a1d24;
    --border-color: #2a2a2a;
    --shadow-color: rgba(0,0,0,0.2);
    --title-color: #f0f0f0;
    --chapter-border: #3a3a3a;
}

.arabic-content {
    font-family: 'Amiri', serif !important;
    font-size: 16px !important;
    text-align: right;
    direction: rtl;
    line-height: 1.4;
    color: var(--text-color);
}

.page-container {
    border: 1px solid var(--border-color);
    padding: 30px;
    border-radius: 8px;
    min-height: 500px;
    background-color: var(--secondary-background);
    box-shadow: 0 2px 4px var(--shadow-color);
}

.title-style {
    font-family: 'Amiri', serif;
    font-size: 20px;
    text-align: center;
    margin-bottom: 2px;
    color: var(--title-color);
}

.chapter-title {
    font-family: 'Amiri', serif;
    font-size: 20px;
    color: var(--title-color);
    text-align: center;
    margin: 12px 0;
    padding-bottom: 5px;
    border-bottom: 1px solid var(--chapter-border);
}

.nav-button {
    margin: 5px;
}

/* Button styling */
.stButton>button {
    border: 1px solid var(--border-color) !important;
}




""", unsafe_allow_html=True)

# MongoDB connection (unchanged)
MONGO_URI = "mongodb+srv://ullah:asad1234@cluster0.572ay.mongodb.net/"
DB_NAME = "pdf_processing"
PDF_COLLECTION = "pdf_metadata"
PAGES_COLLECTION = "pdf_pages"

@st.cache_resource
def init_connection():
    client = MongoClient(MONGO_URI)
    db = client[DB_NAME]
    return db

db = init_connection()

def get_all_books():
    """Get list of all unique PDF books with metadata"""
    return list(db[PDF_COLLECTION].find({}))

def get_pages_for_book(pdf_name):
    """Get all pages for a specific book, sorted by pdf_page_number"""
    return list(db[PAGES_COLLECTION].find({"pdf_name": pdf_name}).sort("pdf_page_number", 1))

def get_chapter_boundaries(book_metadata, pages):
    """Get chapter boundaries based on metadata and actual pages.
    Skips chapters with null/None page numbers and moves to the next chapter."""
    chapters = book_metadata.get("chapters", [])
    if not chapters:
        return [{"name": "Full Book", "start_page": 0, "end_page": len(pages)-1}]
    
    chapter_pages = []
    valid_chapters = []
    
    # First filter out chapters with null/None page numbers
    for chapter in chapters:
        if chapter.get("page_number") is not None:
            valid_chapters.append(chapter)
    
    # If no valid chapters remain after filtering, return full book
    if not valid_chapters:
        return [{"name": "Full Book", "start_page": 0, "end_page": len(pages)-1}]
    
    # Process only valid chapters
    for i, chapter in enumerate(valid_chapters):
        try:
            # Find the actual page index that matches this chapter's page number
            start_page = next(
                idx for idx, p in enumerate(pages) 
                if p["page_number"] == chapter["page_number"]
            )
            
            if i < len(valid_chapters) - 1:
                try:
                    end_page = next(
                        idx for idx, p in enumerate(pages) 
                        if p["page_number"] == valid_chapters[i+1]["page_number"]
                    ) - 1
                except StopIteration:
                    # If next chapter's page not found, use end of book
                    end_page = len(pages) - 1
            else:
                end_page = len(pages) - 1
            
            # Ensure valid page range
            if start_page <= end_page:
                chapter_pages.append({
                    "name": chapter["name"],
                    "start_page": start_page,
                    "end_page": end_page
                })
                
        except StopIteration:
            # Skip this chapter if its page number isn't found in the pages
            continue
    
    # If no valid chapters were processed (all had page numbers not found in pages)
    if not chapter_pages:
        return [{"name": "Full Book", "start_page": 0, "end_page": len(pages)-1}]
    
    # Add "Full Book" option at the beginning
    chapter_pages.insert(0, {
        "name": "Full Book",
        "start_page": 0,
        "end_page": len(pages)-1
    })
    
    return chapter_pages

def main():
    # Initialize session state (unchanged)
    if 'selected_book' not in st.session_state:
        st.session_state.selected_book = None
    if 'current_page' not in st.session_state:
        st.session_state.current_page = 0
    if 'current_chapter' not in st.session_state:
        st.session_state.current_chapter = 0
    if 'view_mode' not in st.session_state:
        st.session_state.view_mode = "chapter"  # or "full"

    # Get all available books
    all_books = get_all_books()
    book_names = [book["pdf_name"] for book in all_books]

    # Sidebar controls
    with st.sidebar:
        st.header("Navigation")
        
        # Book selection dropdown
        selected_book_name = st.selectbox(
            "Select a Book",
            options=book_names,
            index=0 if not st.session_state.selected_book else book_names.index(st.session_state.selected_book["pdf_name"])
        )

        # Update selected book if changed
        selected_book = next((book for book in all_books if book["pdf_name"] == selected_book_name), None)
        if selected_book != st.session_state.selected_book:
            st.session_state.selected_book = selected_book
            st.session_state.current_page = 0
            st.session_state.current_chapter = 0
            st.rerun()

        # Get pages and chapters for selected book
        if st.session_state.selected_book:
            book_pages = get_pages_for_book(st.session_state.selected_book["pdf_name"])
            
            if not book_pages:
                st.warning("No pages found for this book")
                return
            
            chapters = get_chapter_boundaries(st.session_state.selected_book, book_pages)
            
            # View mode selection
            view_mode = st.radio(
                "View Mode",
                ["By Chapter", "Full Book"],
                index=0 if st.session_state.view_mode == "chapter" else 1
            )
            st.session_state.view_mode = "chapter" if view_mode == "By Chapter" else "full"
            
            if st.session_state.view_mode == "chapter":
                # Chapter selection
                chapter_names = [chap["name"] for chap in chapters]
                selected_chapter = st.selectbox(
                    "Select Chapter",
                    options=chapter_names,
                    index=st.session_state.current_chapter
                )
                
                # Update chapter if changed
                new_chapter_idx = chapter_names.index(selected_chapter)
                if new_chapter_idx != st.session_state.current_chapter:
                    st.session_state.current_chapter = new_chapter_idx
                    st.session_state.current_page = chapters[new_chapter_idx]["start_page"]
                    st.rerun()

                # Page navigation within chapter
                current_chapter = chapters[st.session_state.current_chapter]
                total_pages_in_chapter = current_chapter["end_page"] - current_chapter["start_page"] + 1
                
                if total_pages_in_chapter < 1:
                    st.warning("This chapter has no pages")
                    return
                    
                current_chapter_page = st.session_state.current_page - current_chapter["start_page"] + 1
                current_chapter_page = max(1, min(current_chapter_page, total_pages_in_chapter))
                
                page_number = st.number_input(
                    f"Page in chapter (1-{total_pages_in_chapter})",
                    min_value=1,
                    max_value=total_pages_in_chapter,
                    value=current_chapter_page
                )
                
                new_page = current_chapter["start_page"] + page_number - 1
                if new_page != st.session_state.current_page:
                    st.session_state.current_page = new_page
                    st.rerun()
            else:
                # Full book page navigation
                page_number = st.number_input(
                    f"Page in book (1-{len(book_pages)})",
                    min_value=1,
                    max_value=len(book_pages),
                    value=st.session_state.current_page + 1
                )
                
                if page_number - 1 != st.session_state.current_page:
                    st.session_state.current_page = page_number - 1
                    st.rerun()

            # Navigation buttons
            col1, col2 = st.columns(2)
            with col1:
                if st.button("⏮ Previous", use_container_width=True, key="prev"):
                    if st.session_state.current_page > 0:
                        st.session_state.current_page -= 1
                        st.rerun()
            
            with col2:
                if st.button("Next ⏭", use_container_width=True, key="next"):
                    if st.session_state.current_page < len(book_pages) - 1:
                        st.session_state.current_page += 1
                        st.rerun()

            # Book info
            st.header("Book Info")
            st.write(f"**Title:** {st.session_state.selected_book.get('title', '')}")
            st.write(f"**Author:** {st.session_state.selected_book.get('author', '')}")
            
            if st.session_state.view_mode == "chapter":
                current_chapter = chapters[st.session_state.current_chapter]
                st.write(f"**Current Chapter:** {current_chapter['name']}")

    # Main content area
    if st.session_state.selected_book:
        book_pages = get_pages_for_book(st.session_state.selected_book["pdf_name"])
        
        if not book_pages:
            st.warning("No pages found for this book")
            return
            
        if st.session_state.current_page >= len(book_pages):
            st.session_state.current_page = len(book_pages) - 1
            st.rerun()
            
        current_page_data = book_pages[st.session_state.current_page]
        chapters = get_chapter_boundaries(st.session_state.selected_book, book_pages)
        
        # Center the content
        col1, col2, col3 = st.columns([1, 8, 1])
        
        with col2:
            # Display book title
            st.markdown(f'<div class="title-style">{st.session_state.selected_book.get("title", "")}</div>', unsafe_allow_html=True)
            
            # Display chapter title if in chapter mode
            if st.session_state.view_mode == "chapter":
                current_chapter = chapters[st.session_state.current_chapter]
                st.markdown(f'<div class="chapter-title">{current_chapter["name"]}</div>', unsafe_allow_html=True)
            
            # Display the Arabic content
            html_content = current_page_data.get("text", "")
            st.markdown(
                f"""
                <div class="page-container arabic-content">
                {html_content}
                </div>
                """,
                unsafe_allow_html=True
            )
            # Page info
            col1, col2 = st.columns(2)
            with col1:
                st.write(f"**Book Page:** {current_page_data.get('pdf_page_number', 'N/A')}")
            with col2:
                if st.session_state.view_mode == "chapter":
                    current_chapter = chapters[st.session_state.current_chapter]
                    st.write(f"**Chapter Page:** {st.session_state.current_page - current_chapter['start_page'] + 1} of {current_chapter['end_page'] - current_chapter['start_page'] + 1}")
                else:
                    st.write(f"**Page:** {st.session_state.current_page + 1} of {len(book_pages)}")

        # Action buttons section
        st.markdown("---")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("📄 Download Current Page (DOCX)"):
                with st.spinner("Generating DOCX..."):
                    doc = Document()
                    html_to_docx(current_page_data.get("text", ""), doc)
                    docx_buffer = BytesIO()
                    doc.save(docx_buffer)
                    docx_buffer.seek(0)
                    
                    st.download_button(
                        label="⬇️ Save DOCX",
                        data=docx_buffer,
                        file_name=f"page_{current_page_data['pdf_page_number']}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        
        with col2:
            if st.button("📚 Export Full Book (DOCX)"):
                try:
                    with st.spinner("Compiling full book..."):
                        all_pages = get_pages_for_book(st.session_state.selected_book["pdf_name"])
                        
                        if not all_pages:
                            st.error("No pages found for this book")
                            return
                        
                        doc = Document()
                        
                        if st.session_state.selected_book.get("title"):
                            title_para = doc.add_paragraph()
                            title_para.add_run(st.session_state.selected_book["title"]).bold = True
                            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        for i, page in enumerate(all_pages):
                            if not page or not page.get("text"):
                                continue
                            
                            html_to_docx(page["text"], doc)
                            
                            if i % 10 == 0:
                                st.spinner(f"Processing page {i+1}/{len(all_pages)}...")
                        
                        docx_buffer = BytesIO()
                        doc.save(docx_buffer)
                        docx_buffer.seek(0)
                        
                        st.success("Book compilation complete!")
                        st.download_button(
                            label="⬇️ Save Full Book",
                            data=docx_buffer,
                            file_name=f"{st.session_state.selected_book.get('title', 'book')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                
                except Exception as e:
                    st.error(f"Failed to export book: {str(e)}")
        
        with col3:
            if st.button("🔊 Listen to This Page"):
                with st.spinner("Generating audio (may take 20-30 seconds)..."):
                    try:
                        text_content = BeautifulSoup(current_page_data.get("text", ""), "html.parser").get_text()
                        audio_bytes = tts.generate_speech(text_content)

                        st.success("Audio generated!")
                        st.audio(audio_bytes, format="audio/mp3")

                        st.download_button(
                            label="⬇️ Download Audio",
                            data=audio_bytes,
                            file_name=f"page_{current_page_data['pdf_page_number']}.mp3",
                            mime="audio/mp3"
                        )
                    except RuntimeError as e:
                        st.error(f"Audio generation failed. The text might be too long or contain unsupported characters.")
                        st.code(f"Technical details: {str(e)}", language="text")

if __name__ == "__main__":
    main()