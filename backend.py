import os
import re
import fitz  # PyMuPDF
import subprocess
import json
from pymongo import MongoClient
from datetime import datetime
import google.generativeai as genai
from bs4 import BeautifulSoup
import pyarabic.trans
import tempfile
import uuid
from typing import Dict, List, Optional, Tuple
import time
import pyarabic.trans
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
import time
from docx.enum.text import WD_BREAK
# MongoDB Configuration
MONGO_URI = "mongodb+srv://ullah:asad1234@cluster0.572ay.mongodb.net/"
DB_NAME = "pdf_processing"
PDF_COLLECTION = "pdf_metadata"
PAGES_COLLECTION = "pdf_pages"
TEXT_EXTRACTION_PROMPT = """
1. Extract the Arabic text from the provided image and format it in an HTML page. The text should be structured as follows:

    1. All text content should be in `<p>` tags, aligned to the right using `style="text-align: right; direction: rtl;"`. 
    2. Do not use any heading tags (`<h1>`, `<h2>`, etc.) - convert all headings to regular paragraphs.
    3. Do not use bold tags (`<b>`) - keep all text in regular weight.
    4. Any footnotes should be formatted as an ordered list `<ol>` at the end of the page, with each list item `<li>` aligned to the right using `style="text-align: right; direction: rtl;"`.
    6. Do not use a `<style>` tag or external CSS. Apply all styles inline using the `style` attribute.
    7. Do not break paragraphs into multiple `<p>` tags for each line. Use a single `<p>` tag for the entire paragraph.
    8. Donot Extract Header/Footer text.
    9. Do not include any HTML tags other than `<p>`, `<ol>`, and `<li>`.
    Return only the HTML content without additional explanations.

2. Extract 5-10 important keywords/phrases for search purposes:
     - Prioritize proper nouns, technical terms, and repeated concepts
     - Include both single-word and multi-word phrases
     - Exclude very common words (the, and, etc.)
     - List in order of importance
     - Provide in original Arabic script

3. Extract the page number following these rules:
     - Identify page numbers in any common format (Arabic, Roman numerals, or text)
     - Convert all page numbers to numeric format (e.g., '٣' → 3, 'XV' → 15)
     - If no page number exists or is unreadable, use 0
     - Search footer/header/margins if no central page number exists

Return JSON format exactly as shown below - do not include any additional commentary or formatting:

{
    \"text\": \"<p style='text-align: right; direction: rtl;'>...</p><p style='text-align: right; direction: rtl;'>...</p>\",
    \"keywords\": [\"المصطلح الفني\", \"اسم شخص\", \"موضوع رئيسي\"],
    \"page_number\": 3
}
"""

client = MongoClient(MONGO_URI)
db = client[DB_NAME]
pdf_collection = db[PDF_COLLECTION]
pages_collection = db[PAGES_COLLECTION]

class PDFProcessor:
    client = MongoClient(MONGO_URI)
    db = client[DB_NAME]
    pdf_collection = db[PDF_COLLECTION]
    pages_collection = db[PAGES_COLLECTION]

    def __init__(self):
        self.current_book_id = None
        self.gemini_api_keys = [
            "AIzaSyC_FoccwuGrjEMvVJlqq1i7d_Y0ifLqCPw",
            "AIzaSyCkamxFNCtMfhJagHqWctB_Kztt3Or8AK0",
            "AIzaSyCpPsbTaLV9dya1iG0E_PbgmPiCA94CeUo",
            "AIzaSyBfhmN4Rtj69O2sPgkCqPO41xj-BOinaT4",
        "AIzaSyAba4R_F3LyYOajY2hF8_qBFy3ucAsanmk"        ]
        self.current_key_idx = 0
        self.temp_image_dir = "temp_pdf_images"
    def remove_small_number_brackets(self,input_string:str):
        # Regular expression to match brackets containing one or two digits (English or Arabic) with optional spaces
        digit_text=pyarabic.trans.normalize_digits(input_string, source='all', out='west')

        cleaned_string = re.sub(r"\(\d+\)", "", digit_text)
        return cleaned_string


    def remove_square_brackets(self,input_string:str):
        cleaned_text = re.sub(r"\[[\u0600-\u06FF\s\d/]+\]", "", input_string)
        return cleaned_text
    def generate_book_id(self) -> str:
        return str(uuid.uuid4())

    def clean_pdf_name(self, pdf_path: str) -> str:
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]
        return re.sub(r'[\d()\[\]]', '', base_name).strip()

    def check_pdf_exists(self, pdf_name: str) -> bool:
        return pdf_collection.count_documents({"pdf_name": pdf_name}) > 0

    def get_existing_pages(self, book_id: str) -> set:
        existing = pages_collection.find({"book_id": book_id}, {"pdf_page_number": 1})
        return {doc['pdf_page_number'] for doc in existing}

    def extract_metadata_with_gemini(self, pdf_path: str) -> Dict:
        try:
            genai.configure(api_key=self.gemini_api_keys[self.current_key_idx])
            model = genai.GenerativeModel("gemini-2.0-flash")
            
            with tempfile.NamedTemporaryFile(suffix=".pdf") as temp_pdf:
                src = fitz.open(pdf_path)
                dst = fitz.open()
                
                total_pages = len(src)
                pages_to_extract = list(range(5)) + list(range(max(5, total_pages-5), total_pages))
                
                for page_num in pages_to_extract:
                    dst.insert_pdf(src, from_page=page_num, to_page=page_num)
                
                dst.save(temp_pdf.name)
                src.close()
                dst.close()

                myfile = genai.upload_file(temp_pdf.name)
                prompt = """
                Analyze this Arabic book/document PDF and extract the following metadata in STRICT JSON format:
                {
                    "title": "Book title in Arabic",
                    "author": "Author name in Arabic",
                    "subject": "Subject/topic in Arabic",
                    "chapters": [
                        {
                            "name": "Chapter name in Arabic",
                            "page_number": X
                        },
                        ... (all chapters)
                    ]

                }

                Rules:
                1. Extract all information in the original Arabic text
                2. Include accurate page numbers for each chapter
                3. If any field is not available, set it to an empty string
                4. Ensure the output is valid JSON that can be parsed directly
                5. donot include the diacritics

                """

                response = model.generate_content([prompt, myfile])
                response_text = response.text
                
                start_idx = response_text.find('{')
                end_idx = response_text.rfind('}') + 1
                metadata = json.loads(response_text[start_idx:end_idx])
                
                genai.delete_file(myfile.name)
                return metadata

        except Exception as e:
            print(f"Metadata extraction error: {str(e)}")
            self.current_key_idx = (self.current_key_idx + 1) % len(self.gemini_api_keys)
            return None
    def convert_english_to_arabic_digits(self,text:str):
        """Convert English digits to Arabic digits"""
        digit_mapping = {
            '0': '٠', '1': '١', '2': '٢', '3': '٣', '4': '٤',
            '5': '٥', '6': '٦', '7': '٧', '8': '٨', '9': '٩'
        }
        return ''.join(digit_mapping.get(c, c) for c in text)

    def fix_inverted_brackets(self,text:str):
        """Fix inverted brackets in RTL text"""
        fixed_text = []
        for char in text:
            if char in '()[]{}':
                fixed_text.append('\u200E' + char + '\u200E')
            else:
                fixed_text.append(char)
        return ''.join(fixed_text)

    def pdf_to_images(self, pdf_path: str, page_range: range = None) -> List[str]:
        """Convert PDF pages to images"""
        if not os.path.exists(self.temp_image_dir):
            os.makedirs(self.temp_image_dir)

        pdf = fitz.open(pdf_path)
        image_paths = []

        for page_num in (page_range if page_range else range(len(pdf))):
            page = pdf.load_page(page_num)
            pix = page.get_pixmap(dpi=300)  # Higher DPI for better OCR
            img_path = os.path.join(self.temp_image_dir, f"{page_num+1}.jpg")
            pix.save(img_path)
            image_paths.append(img_path)

        pdf.close()
        return image_paths

    def process_page_image(self, image_path: str, pdf_name: str) -> Dict:
        """Process single page image with Gemini"""
        try:
            genai.configure(api_key=self.gemini_api_keys[self.current_key_idx])
            model = genai.GenerativeModel("models/gemini-2.0-flash")
            
            print(f"Processing {image_path}")
            myfile = genai.upload_file(image_path)
            result = model.generate_content([myfile, TEXT_EXTRACTION_PROMPT])

            # Extract JSON from response
            response = result.text
            json_str = response[response.find('{'):response.rfind('}')+1]
            page_data = json.loads(json_str)
            
            # Get page number from filename
            pdf_page_number = int(os.path.basename(image_path).split('.')[0])
            print(pdf_page_number)
            # Clean and process text
            soup = BeautifulSoup(page_data['text'], 'html.parser')

            # Remove all <ol> and <li> tags and their content
            for ol_tag in soup.find_all('ol'):
                ol_tag.decompose()
            for li_tag in soup.find_all('li'):
                li_tag.decompose()

            # Process remaining text content
            for elem in soup.find_all(text=True):
                text=elem
                text=self.remove_square_brackets(text)
                text=self.remove_small_number_brackets(text)
                elem.replace_with(self.fix_inverted_brackets(self.convert_english_to_arabic_digits(text)))

            return {
                'pdf_name': pdf_name,
                'page_number': int(page_data['page_number']),
                'text': str(soup),
                'keywords': page_data.get('keywords', []),
                'processed_at': datetime.now(),
                'image_path': image_path,
                'pdf_page_number': pdf_page_number,
                'original_response': page_data  # Still preserve original response
            }, self.current_key_idx
        
        except Exception as e:
            print(f"Page processing error: {str(e)}")
            self.current_key_idx = (self.current_key_idx + 1) % len(self.gemini_api_keys)
            return None, self.current_key_idx
        finally:
            if 'myfile' in locals():
                try:
                    genai.delete_file(myfile.name)
                except:
                    pass

    def store_page_data(self, book_id: str, page_data: Dict):
        """Store processed page data with all original attributes"""
        page_data["book_id"] = book_id
        pages_collection.update_one(
            {"book_id": book_id, "pdf_page_number": page_data["pdf_page_number"]},
            {"$set": page_data},
            upsert=True
        )

    def process_pdf_pages(
        self,
        pdf_path: str,
        pdf_name:str,
        start_page: int = 1,
        end_page: int = None,
        book_id: str = None
    ) -> Tuple[str, Dict]:
        """Process PDF pages as images with Gemini"""
        # Initialize or get existing book record
        if book_id is None:
            if self.check_pdf_exists(pdf_name):
                existing = pdf_collection.find_one({"pdf_name": pdf_name})
                book_id = existing["book_id"]
                metadata = existing
            else:
                book_id = self.generate_book_id()
                metadata = self.extract_metadata_with_gemini(pdf_path)
                if not metadata:
                    return None, None
                
                pdf_collection.insert_one({
                    "book_id": book_id,
                    "pdf_name": pdf_name,
                    "title": metadata.get("title", ""),
                    "author": metadata.get("author", ""),
                    "subject": metadata.get("subject", ""),
                    "chapters": metadata.get("chapters", []),
                    "created_at": datetime.now()
                })
        else:
            metadata = pdf_collection.find_one({"book_id": book_id})

        # Convert PDF to images
        doc = fitz.open(pdf_path)
        total_pages = len(doc)
        end_page = end_page if end_page else total_pages
        existing_pages = self.get_existing_pages(book_id)
        doc.close()

        # Process pages
        processing_report = {
            "total_pages": total_pages,
            "processed_pages": list(existing_pages),
            "newly_processed": [],
            "skipped_pages": []
        }

        # Convert only the needed pages to images
        page_range = range(start_page-1, min(end_page, total_pages))
        image_paths = self.pdf_to_images(pdf_path, page_range)

        for img_path in image_paths:
            page_num = int(os.path.basename(img_path).split('.')[0])
            
            if page_num in existing_pages:
                processing_report["skipped_pages"].append(page_num)
                continue
            
            page_data, key_idx = self.process_page_image(img_path, pdf_name)
            self.current_key_idx = key_idx
            
            if page_data:
                self.store_page_data(book_id, page_data)
                processing_report["newly_processed"].append(page_num)
                processing_report["processed_pages"].append(page_num)
            
            

        # Cleanup images
        for img_path in image_paths:
            try:
                os.remove(img_path)
            except:
                pass

        return book_id, processing_report

class TextToSpeech:
    client = MongoClient(MONGO_URI)
    db = client[DB_NAME]
    pdf_collection = db[PDF_COLLECTION]
    pages_collection = db[PAGES_COLLECTION]

    @staticmethod
    def generate_speech(text: str, voice: str = "ar-AE-HamdanNeural") -> bytes:
        """Robust Arabic audio generation with error handling"""
        try:
            # Clean problematic characters
            cleaned_text = text.replace('"', "'").replace('\n', ' ')

            with tempfile.NamedTemporaryFile(suffix=".mp3", delete=False) as temp_audio:
                temp_path = temp_audio.name

            # Split long text into chunks (edge-tts has ~5000 char limit)
            max_chunk = 3000
            chunks = [cleaned_text[i:i+max_chunk] for i in range(0, len(cleaned_text), max_chunk)]

            audio_files = []
            for i, chunk in enumerate(chunks):
                chunk_path = f"{temp_path}_{i}.mp3"
                cmd = [
                    "edge-tts",
                    "--voice", voice,
                    "--text", chunk,
                    "--write-media", chunk_path
                ]

                # Run with timeout
                try:
                    subprocess.run(cmd, check=True, timeout=30)
                    audio_files.append(chunk_path)
                except subprocess.TimeoutExpired:
                    print(f"Timeout for chunk {i}")
                    continue

            # Combine chunks if multiple
            if len(audio_files) > 1:
                from pydub import AudioSegment
                combined = AudioSegment.empty()
                for f in audio_files:
                    combined += AudioSegment.from_mp3(f)
                combined.export(temp_path, format="mp3")
            elif audio_files:
                import shutil
                shutil.move(audio_files[0], temp_path)

            # Read final audio
            with open(temp_path, "rb") as f:
                return f.read()

        except Exception as e:
            print(f"Audio generation error: {str(e)}")
            raise RuntimeError(f"Couldn't generate audio: {str(e)}")
        finally:
            # Cleanup
            import os
            if 'temp_path' in locals() and os.path.exists(temp_path):
                os.unlink(temp_path)
            for f in audio_files:
                if os.path.exists(f):
                    os.unlink(f)
class DatabaseManager:
    
    client = MongoClient(MONGO_URI)
    db = client[DB_NAME]
    pdf_collection = db[PDF_COLLECTION]
    pages_collection = db[PAGES_COLLECTION]
    @staticmethod
    def get_processing_status(book_id: str) -> Dict:
        """Get processing status for all pages of a book"""
        total_pages = pdf_collection.find_one(
            {"book_id": book_id}, 
            {"total_pages": 1}
        ).get("total_pages", 0)
        
        processed_pages = list(pages_collection.find(
            {"book_id": book_id},
            {"page_number": 1}
        ).sort("page_number", 1))
        
        processed_numbers = [p["page_number"] for p in processed_pages]
        all_pages = set(range(1, total_pages + 1))
        
        return {
            "total_pages": total_pages,
            "processed_pages": processed_numbers,
            "unprocessed_pages": list(all_pages - set(processed_numbers))
        }

    @staticmethod
    def get_book(book_id: str) -> Optional[Dict]:
        return pdf_collection.find_one({"book_id": book_id})
    
    @staticmethod
    def get_page(book_id: str, page_number: int) -> Optional[Dict]:
        return pages_collection.find_one({
            "book_id": book_id,
            "page_number": page_number
        })
    
    @staticmethod
    def search_books(query: str) -> List[Dict]:
        regex = re.compile(query, re.IGNORECASE)
        return list(pdf_collection.find({
            "$or": [
                {"title": regex},
                {"author": regex},
                {"subject": regex}
            ]
        }))
def convert_english_to_arabic_digits(text):
    # Mapping of English digits to Arabic digits
    digit_mapping = {
        '0': '٠',
        '1': '١',
        '2': '٢',
        '3': '٣',
        '4': '٤',
        '5': '٥',
        '6': '٦',
        '7': '٧',
        '8': '٨',
        '9': '٩'
    }

    # Replace each English digit with its corresponding Arabic digit
    for eng, arb in digit_mapping.items():
        text = text.replace(eng, arb)

    return text
def set_paragraph_direction(paragraph, direction='rtl'):
    """
    Set the text direction of a paragraph to RTL or LTR.
    """
    p_pr = paragraph._element.get_or_add_pPr()
    p_bidi = OxmlElement('w:bidi')
    p_bidi.set(qn('w:val'), '1' if direction == 'rtl' else '0')
    p_pr.append(p_bidi)

def fix_inverted_brackets(text):
    """
    Fix inverted brackets in RTL text by adding Unicode control characters.
    """
    # Characters that should not be mirrored in RTL text
    fixed_text = []
    for char in text:
        if char in '()[]{}':
            # Add Unicode LEFT-TO-RIGHT MARK (LRM) before and after the character
            fixed_text.append('\u200E' + char + '\u200E')
        else:
            fixed_text.append(char)
    return ''.join(fixed_text)

def extract_keywords_from_html(html_content):
    """
    Extract headings and bold words from HTML content.
    """
    soup = BeautifulSoup(html_content, 'html.parser')
    keywords = []

    # Extract headings
    for heading in soup.find_all('h1'):
        keywords.append(heading.get_text().strip())

    # Extract bold words
    for bold in soup.find_all('b'):
        keywords.append(bold.get_text().strip())

    return list(set(keywords))  # Remove duplicates

def html_to_docx(html_content, doc):
    """
    Convert HTML content to DOCX format while preserving formatting and alignment.
    """
    soup = BeautifulSoup(html_content, 'html.parser')

    for element in soup.find_all(['h1', 'p', 'ol', 'li']):
        if element.name == 'ol':
            # Add ordered list
            list_number = 1  # Reset list numbering for each new page
            for li in element.find_all('li'):
                # Add a paragraph without the 'List Number' style
                paragraph = doc.add_paragraph()
                set_paragraph_direction(paragraph, 'rtl')  # Set RTL for Arabic text

                # Add the list number manually (convert to Arabic digits)
                arabic_number = convert_english_to_arabic_digits(str(list_number))
                paragraph.add_run(f" {arabic_number}. ").bold = True
                list_number += 1  # Increment the list number

                # Add the list item content
                for content in li.contents:
                    if content.name == 'b':
                        text = convert_english_to_arabic_digits(fix_inverted_brackets(content.get_text()))
                        run = paragraph.add_run(text)
                        run.font.size = Pt(13)  # 18 pt font size

# Set the font name
                        run.font.name = 'Cambria'

                    elif content.name == 'sup':
                        # Add the sup content without the <sup> tags
                        text = convert_english_to_arabic_digits(fix_inverted_brackets(content.get_text()))
                        run=paragraph.add_run(text)
                        run.font.size = Pt(13)  # 18 pt font size

# Set the font name
                        run.font.name = 'Cambria'
                    else:
                        text = convert_english_to_arabic_digits(fix_inverted_brackets(str(content)))
                        run=paragraph.add_run(text)
                        run.font.size = Pt(13)  # 18 pt font size

# Set the font name
                        run.font.name = 'Cambria'

        elif element.name == "p" or element.name == "h1":
            # Add paragraph
            paragraph = doc.add_paragraph()
            set_paragraph_direction(paragraph, 'rtl')  # Set RTL for Arabic text

            # Handle content
            for content in element.contents:
                if content.name == 'b':
                    text = convert_english_to_arabic_digits(fix_inverted_brackets(content.get_text()))
                    run = paragraph.add_run(text)
                    run.font.size = Pt(13)  # 18 pt font size

# Set the font name
                    run.font.name = 'Cambria'
                elif content.name == 'sup':
                    # Add the sup content without the <sup> tags
                    text = convert_english_to_arabic_digits(fix_inverted_brackets(content.get_text()))
                    run=paragraph.add_run(text)
                    run.font.size = Pt(13)  # 18 pt font size

# Set the font name
                    run.font.name = 'Cambria'
                elif content.name == 'br':
                    # Add the sup content without the <sup> tags
                    text = convert_english_to_arabic_digits(fix_inverted_brackets(content.get_text()))
                    run=paragraph.add_run(text)
                    run.font.size = Pt(13)  # 18 pt font size

# Set the font name
                    run.font.name = 'Cambria'
                elif content.name == 'span':
                    # Add the sup content without the <sup> tags
                    text = convert_english_to_arabic_digits(fix_inverted_brackets(content.get_text()))
                    run=paragraph.add_run(text)
                    run.font.size = Pt(13)  # 18 pt font size

# Set the font name
                    run.font.name = 'Cambria'
                else:
                    text = convert_english_to_arabic_digits(fix_inverted_brackets(str(content)))
                    run=paragraph.add_run(text)
                    run.font.size = Pt(13)  # 18 pt font size
# Set the font name
                    run.font.name = 'Cambria'

# Add a page break after each page  # Add a page break after each page
    paragraph = doc.add_paragraph()
    paragraph.add_run("-------------------------------------------------------------")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
